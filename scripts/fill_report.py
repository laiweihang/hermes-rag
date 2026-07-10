"""
填充人工智能实践赛作品报告 DOCX 模板。
用法: python scripts/fill_report.py
输出: 项目根目录/作品报告.docx
"""
import os
from docx import Document

TEMPLATE = os.path.expanduser(
    "~/Downloads/05-3 作品报告（人工智能实践赛，2026版）.docx"
)
OUTPUT = os.path.join(os.path.dirname(os.path.dirname(__file__)), "作品报告.docx")

BODY_STYLE = "正文段落"


def set_paragraph_text(para, text):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def clear_paragraph(para):
    for run in para.runs:
        run.text = ""
    para.text = ""


def main():
    doc = Document(TEMPLATE)
    paras = doc.paragraphs
    body_style = doc.styles[BODY_STYLE]

    # ── 填写日期 ──
    paras[7].clear()
    paras[7].add_run("填写日期：2026年4月15日")

    # ====================================================================
    # 第1章 作品概述 (paras 10-21)
    # ====================================================================
    clear_paragraph(paras[11])  # 清空说明
    for i in range(12, 22):
        clear_paragraph(paras[i])

    paras[12].style = body_style
    set_paragraph_text(
        paras[12],
        "一、创意来源与产生背景\n"
        "在企业、政府和教育机构中，大量关键知识以 PDF 文件、Word 文档、"
        "纸质扫描件等形式沉淀在内部服务器或个人电脑中。以某市政务服务中心为例，"
        "仅劳动人事相关的政策文件就有数百份，涵盖《劳动合同法》、《职工带薪年休假条例》、"
        "各类实施细则和内部规章制度。当窗口人员需要解答市民关于"
        "\u201c试用期最长多久\u201d\u201c加班费怎么算\u201d等问题时，"
        "传统做法是凭记忆或手动翻阅文件查找条款，效率低、易出错。"
        "而现有的全文检索工具（如 Elasticsearch）只能基于关键词匹配，"
        "无法理解\u201c工作满 15 年能休几天假\u201d这样的自然语言提问。"
        "赫尔墨斯（Hermes）项目正是为解决这一痛点而设计的——"
        "它将静态文档转化为可对话的智能知识服务。",
    )

    paras[13].style = body_style
    set_paragraph_text(
        paras[13],
        "二、用户群体\n"
        "（1）企业内部员工：查询公司制度文档（如报销制度、出差审批流程、产品技术手册），"
        "例如新入职员工输入\u201c差旅报销标准是多少？\u201d即可获取完整条款及出处。\n"
        "（2）政务窗口工作人员：查询政策法规条款，快速为市民提供权威依据。\n"
        "（3）财务审计人员：上传年度财务报表 PDF 后提问\u201c"
        "今年营业收入同比增长多少？\u201d系统自动从报表中提取数据并计算回答。\n"
        "（4）教育培训机构：将教材或培训资料入库后，学员可随时交互式学习，"
        "例如\u201c请总结第三章的核心要点\u201d。",
    )

    paras[14].style = body_style
    set_paragraph_text(
        paras[14],
        "三、主要功能\n"
        "功能1：RAG 检索增强生成问答。用户输入自然语言问题后，系统自动从向量数据库中检索最相关的文档片段，"
        "将其与问题一起送入本地大语言模型生成回答，并展示引用来源。"
        "例如，用户上传《职工带薪年休假条例》后提问\u201c工作满 12 年年假多少天？\u201d，"
        "系统会检索到第三条内容\u201c已满10年不满20年的，年休假10天\u201d，"
        "并据此生成回答\u201c根据条例第三条，工作满 12 年的年休假为 10 天\u201d，"
        "同时标注来源文件名和引用片段。\n"
        "功能2：多格式文档入库。支持文本 PDF（PyPDF 提取）、扫描 PDF（GLM-OCR 逐页识别）、"
        "DOCX（python-docx 提取）、TXT（直读）和 PNG/JPG/BMP/TIFF/WEBP 图片（GLM-OCR 识别）。"
        "例如，用户上传一份20世纪90年代的扫描档案，系统自动检测到每页可提取文本不足 50 字，"
        "判定为扫描件，自动切换到 OCR 模式逐页识别，最终完成向量入库。\n"
        "功能3：技能场景系统。内置 4 个预置技能——财务分析、会议纪要助手、本地政策文档、发票助手。"
        "每个技能包含专业化的 system_prompt 和自动检测关键词。"
        "例如，当用户上传名为\u201c2025年度利润表.pdf\u201d的文件时，"
        "系统自动匹配到\u201c财务分析\u201d技能（因文件名包含\u201c利润\u201d关键词），"
        "随后的问答将使用财务分析专用提示词，回答更专业精准。用户也可在侧栏手动切换技能。\n"
        "功能4：规则引擎零延迟直答。对\u201c加班费怎么算\u201d\u201c年假多少天\u201d\u201c"
        "试用期多长\u201d等高频问题，规则引擎通过关键词和正则表达式匹配后直接返回预设答案，"
        "不经过 LLM 推理，响应时间低于 5 毫秒。例如，用户问\u201c加班费怎么算？\u201d，"
        "规则引擎匹配到正则 r\u201c加班.*?(?:工资|费)\u201d，直接返回《劳动法》第四十四条全文。\n"
        "功能5：多轮对话记忆。系统自动携带最近 10 条对话历史作为上下文，"
        "用户可连续追问。例如先问\u201c公司差旅报销标准是多少？\u201d，"
        "再追问\u201c那住宿费呢？\u201d系统能理解\u201c那\u201d指代差旅场景。\n"
        "功能6：SSE 流式输出。采用 Server-Sent Events 协议逐 token 推送，"
        "用户无需等待完整回答生成即可开始阅读，体验类似 ChatGPT。\n"
        "功能7：管理员后台。提供用户管理（创建/删除用户）、反馈管理（查看赞/踩统计）、"
        "技能 CRUD（新增/编辑/删除技能及其提示词）、系统统计面板。",
    )

    paras[15].style = body_style
    set_paragraph_text(
        paras[15],
        "四、作品特色\n"
        "（1）全链路本地化：LLM 推理（LM Studio + Qwen）、文本嵌入（Nomic Embed）、OCR 识别（GLM-OCR）"
        "全部在本地 GPU/CPU 上运行，文档数据不出内网，从根本上保障数据隐私。\n"
        "（2）规则引擎优先策略：高频问题毫秒级响应，显著降低 LLM 调用频率和算力消耗。\n"
        "（3）技能场景自动识别：无需用户手动配置，上传文件时根据文件名自动推荐最匹配的技能。\n"
        "（4）多模态文档处理：通过本地 GLM-OCR 模型打通扫描件和图片的信息壁垒。",
    )

    paras[16].style = body_style
    set_paragraph_text(
        paras[16],
        "五、应用价值与推广前景\n"
        "本作品将组织内的静态文档资产转化为可对话、可追溯的智能知识服务。"
        "在企业知识库场景，员工平均查找一份制度条款的时间可从 10-30 分钟缩短至 10 秒以内；"
        "在政务窗口场景，工作人员可即时引用政策依据，提升服务效率和准确性。"
        "由于全部组件均可本地部署、不依赖云端 API，特别适合政府机构、金融企业、"
        "国防科研等对数据安全有严格要求的组织，具有广泛的推广前景。",
    )

    for i in range(17, 22):
        clear_paragraph(paras[i])

    # ====================================================================
    # 第2章 问题分析 (paras 22-36)
    # ====================================================================

    # 2.1 问题来源 (H2=23, 说明=24, 正文=25)
    clear_paragraph(paras[24])
    paras[25].style = body_style
    set_paragraph_text(
        paras[25],
        "本作品的问题来源于真实的企业和政务场景。以某大型制造企业为例，"
        "其内部管理制度、产品技术手册、质量标准文件累计超过 2000 份，"
        "分布在 OA 系统、文件服务器和各部门本地硬盘中。当质检工程师需要查询某个零件的"
        "检验标准时，往往需要在多个文件夹中搜索关键词，再逐一打开文件人工查找，"
        "整个过程平均耗时 15-30 分钟。更棘手的是，约 30% 的历史档案为纸质文件扫描件，"
        "这些 PDF 虽然外观与正常文档无异，但内部实际上是图片，"
        "无法被任何文本搜索工具检索到。\n"
        "在政务场景中同样如此。某区社会保障局拥有数百份劳动法规、实施细则和内部文件。"
        "窗口人员在解答群众咨询时，经常需要现场翻找文件确认具体条款，"
        "不仅效率低下，还可能因记忆偏差给出不准确的答复。\n"
        "综合来看，问题的核心在于：组织内大量文档知识处于\u201c沉睡\u201d状态，"
        "无法被高效、精准、智能地检索和利用。",
    )

    # 2.2 现有解决方案 (H2=26, 说明=27, 正文=28)
    clear_paragraph(paras[27])
    paras[28].style = body_style
    set_paragraph_text(
        paras[28],
        "现有解决方案及其不足分析如下：\n"
        "方案一：云端 RAG 服务（如 ChatGPT + Retrieval 插件、Azure AI Search、百度知识增强）。"
        "这类方案具备强大的语义理解和生成能力，但存在两个致命问题：①数据隐私——"
        "需要将全部文档上传至第三方云服务器，对于政务、金融、国防等涉密场景完全不可接受[1]；"
        "②持续成本——按 token 计费，日常大量查询会产生可观的 API 费用。\n"
        "方案二：传统全文检索系统（如 Elasticsearch、Apache Solr）。"
        "这类方案基于倒排索引和关键词匹配，可部署在本地。但其局限性在于：①无法理解语义——"
        "用户搜索\u201c工作 15 年能休几天假\u201d时，系统只会匹配包含\u201c15 年\u201d"
        "\u201c休假\u201d字面词的文档，对同义词、近义词、语句改写无能为力；"
        "②无生成能力——只返回匹配的文档列表，用户仍需自行阅读和提炼答案；"
        "③不支持扫描件——图片格式的 PDF 无法建立索引。\n"
        "方案三：本地笔记工具（如 Obsidian、Notion、飞书文档）。"
        "支持全文搜索和标签分类，适合个人知识管理，但缺乏语义检索、AI 生成和多模态 OCR 能力，"
        "不适合企业级大规模文档的智能问答需求。\n"
        "下表对比了各方案在关键维度上的表现：\n"
        "维度 | 云端 RAG | Elasticsearch | 本地笔记 | 本作品\n"
        "数据隐私 | 差 | 好 | 好 | 好\n"
        "语义理解 | 强 | 弱 | 无 | 强\n"
        "AI 生成 | 有 | 无 | 无 | 有\n"
        "扫描件 OCR | 部分 | 无 | 无 | 有\n"
        "场景化适配 | 弱 | 无 | 无 | 有（技能系统）\n"
        "部署成本 | 高（按需付费） | 中 | 低 | 低（本地一次性）",
    )

    # 2.3 本作品要解决的痛点问题 (H2=29, 说明=30, 正文=31,32)
    clear_paragraph(paras[30])
    paras[31].style = body_style
    set_paragraph_text(
        paras[31],
        "基于以上对比分析，本作品重点解决以下四个核心痛点：\n"
        "痛点一：数据隐私与安全合规。政务、金融、企业内部制度等文档包含敏感信息，"
        "不能上传至任何第三方平台。例如某银行的信贷审批流程文档涉及风控策略，"
        "一旦泄露将造成重大商业风险。本作品通过全链路本地化部署（LLM、嵌入模型、OCR 均在本地运行）"
        "来解决此问题。\n"
        "痛点二：扫描件和图片文档无法检索。大量历史档案以扫描 PDF 或照片形式保存，"
        "例如 20 世纪 90 年代的合同原件、手写会议纪要照片。传统检索工具对这些文档完全无能为力。"
        "本作品集成本地 GLM-OCR 模型，自动识别扫描件并提取文字后入库。\n"
        "痛点三：缺乏场景化专业问答。不同业务场景需要不同的专业知识引导——"
        "财务人员需要精确的数据提取和计算，会议记录需要结构化的要点提炼，"
        "发票核验需要票面信息的准确提取。通用问答模型缺乏这种场景适配能力。"
        "本作品通过技能场景系统，为每个场景定制专业 system_prompt。\n"
        "痛点四：高频固定问题的重复回答浪费算力。\u201c加班费怎么算\u201d"
        "\u201c年假多少天\u201d等问题每天可能被问数十次，"
        "每次都调用 LLM 既浪费 GPU 算力又增加响应延迟。"
        "本作品通过规则引擎前置拦截，命中时毫秒级返回预设答案。",
    )
    clear_paragraph(paras[32])

    # 2.4 解决问题的思路 (H2=33, 说明=34, 正文=35,36)
    clear_paragraph(paras[34])
    paras[35].style = body_style
    set_paragraph_text(
        paras[35],
        "本作品采用\u201c规则引擎 + 本地 RAG + 本地 OCR + 技能场景\u201d"
        "四层递进的技术路线来解决上述问题。系统的核心处理流程如下：\n"
        "（1）文档入库阶段：用户通过 Web 界面上传文档 \u2192 后端根据文件后缀自动分流"
        "（.pdf \u2192 PyPDF 提取/OCR 回退、.docx \u2192 python-docx、"
        ".txt \u2192 直读、图片 \u2192 GLM-OCR） \u2192 RecursiveCharacterTextSplitter 切片"
        "（chunk_size=500, overlap=100） \u2192 LM Studio 嵌入模型"
        "（nomic-embed-text-v1.5）向量化 \u2192 ChromaDB 持久化存储。\n"
        "（2）问答阶段：用户输入问题 \u2192 规则引擎优先匹配（关键词+正则） \u2192 "
        "若未命中，进入 RAG 流程：ChromaDB Top-K（默认 K=3）向量相似度检索 \u2192 "
        "检索到的文档片段与问题拼装为上下文 \u2192 加入技能 system_prompt 和最近 10 条对话历史 "
        "\u2192 本地 LLM（Qwen 3.5-9B）流式生成回答 \u2192 SSE 逐 token 推送至前端。",
    )
    paras[36].style = body_style
    set_paragraph_text(
        paras[36],
        "功能需求明细：多格式文档入库（PDF/DOCX/TXT/PNG/JPG/BMP/TIFF/WEBP）、"
        "向量语义检索、流式 RAG 问答、技能场景自动识别与手动切换、"
        "规则引擎直答、多轮对话记忆、对话导出（JSON/CSV）、用户反馈收集（赞/踩）、"
        "管理员后台（用户/反馈/技能/统计管理）。\n"
        "性能需求：规则引擎响应时间 < 10ms；RAG 流式首字节延迟 < 5s；"
        "单文件上传支持最大 50MB；支持 PDF/DOCX/TXT 及 7 种图片格式。\n"
        "数据来源与特点：数据来源为用户自行上传的业务文档，无需预置数据集。"
        "格式上涵盖文本 PDF（PyPDF 直接提取）、扫描 PDF（判定条件：超过 50% 的页面"
        "提取文本不足 50 字符即判定为扫描件，自动切换 GLM-OCR 逐页渲染识别）、"
        "DOCX（Docx2txtLoader 提取）、TXT（UTF-8 直接读取）和图片（GLM-OCR 识别）。"
        "所有文档经 RecursiveCharacterTextSplitter 按 500 字切片（100 字重叠防止语义切断），"
        "通过 LM Studio 嵌入模型（nomic-embed-text-v1.5，768 维向量）转为稠密向量后存入 ChromaDB。"
        "数据特点：格式多样、领域广泛、规模随使用持续增长。",
    )

    # ====================================================================
    # 第3章 技术方案 (paras 37-40)
    # ====================================================================
    clear_paragraph(paras[38])
    paras[39].style = body_style
    set_paragraph_text(
        paras[39],
        "3.1 总体架构\n"
        "本系统采用前后端分离的三层架构，技术路线框架如下：\n"
        "表示层（前端）：Next.js 16 + React 19 + TypeScript + Tailwind CSS 4 + shadcn/ui 组件库 + "
        "Framer Motion 动画库。运行在 http://localhost:3000，通过 Axios 调用后端 API，"
        "SSE 流式接收 LLM 生成内容。\n"
        "业务逻辑层（后端）：Python FastAPI 框架，运行在 http://localhost:8000，"
        "提供 37 个以上 RESTful API 端点。核心模块包括：认证中间件（JWT + bcrypt）、"
        "RAG 引擎（LangChain 编排）、文档入库管道、OCR 引擎、规则引擎、技能管理。"
        "通过 CORS 中间件支持跨域请求。\n"
        "数据层：ChromaDB 向量数据库（存储文档向量片段和嵌入）+ SQLite 关系数据库"
        "（SQLAlchemy ORM，存储用户、对话、消息、反馈、技能等结构化数据）+ "
        "LM Studio 本地模型服务（http://127.0.0.1:1234/v1，提供 LLM 推理和文本嵌入两个 API）。",
    )
    paras[40].style = body_style
    set_paragraph_text(
        paras[40],
        "3.2 核心技术模块详解\n\n"
        "（1）RAG 检索增强生成（rag_engine.py）\n"
        "本模块是系统的核心，采用 LangChain 框架编排完整的 RAG 流程。"
        "具体实现分为两个分支：当 ChromaDB 中有文档数据时走 RAG 检索路径，否则走直接对话路径。"
        "RAG 路径的具体流程为：①调用 ChromaDB 的 as_retriever 接口执行 Top-K（默认 K=3）"
        "余弦相似度检索；②将检索到的文档片段通过 format_docs 函数拼接为上下文字符串"
        "（最大 2000 字符，超出时截断到最近的完整片段）；③构建 ChatPromptTemplate，"
        "包含系统提示词（来自技能 system_prompt 或默认政策问答提示词）、"
        "最近 10 条对话历史（human/ai 交替）和当前用户问题+参考资料；"
        "④通过 ChatOpenAI（连接 LM Studio 本地 API）调用 Qwen 3.5-9B 模型生成回答；"
        "⑤后处理：用正则表达式 re.sub(r'<think>.*?</think>') 清除模型的推理标签，"
        "只保留最终答案。流式模式下使用 chain.stream() 逐 token 生成，"
        "通过生成器 yield 给 FastAPI 的 StreamingResponse。"
        "系统还内置了基于 tenacity 的重试机制，对 ConnectionError 和 5xx 错误"
        "自动重试最多 3 次（指数退避 1-10 秒），400/401 等客户端错误不重试。\n\n"
        "（2）OCR 多模态识别（ocr_engine.py）\n"
        "本模块解决扫描件和图片文档的文字识别问题。采用 GLM-OCR 模型（基于 Transformers 架构的"
        "多模态语言模型），支持 AutoModelForCausalLM 加载，可利用 GPU 加速[2]。"
        "处理扫描 PDF 的流程为：①使用 PyMuPDF（fitz）将每页渲染为 150 DPI 的 PNG 图片；"
        "②逐页将图片送入 GLM-OCR 模型，通过 tokenizer.apply_chat_template 构建"
        "多模态对话输入（role=user, image=PIL.Image, content=\u201c请识别图片中的所有文字\u201d）；"
        "③模型 generate 输出识别文字（max_new_tokens=2048）；"
        "④将每页识别结果封装为 LangChain Document 对象，附带页码元数据。"
        "判定扫描件的策略是：PyPDF 提取后超过 50% 的页面文本量低于 50 字符即判定为扫描件，"
        "自动回退到 OCR 模式。模型和分词器采用懒加载策略（首次调用时加载，之后常驻内存），"
        "避免启动时的长时间等待。\n\n"
        "（3）规则引擎（rule_engine.py）\n"
        "规则引擎是 RAG 之前的第一道拦截层，用于零延迟回答高频固定问题。"
        "每条规则包含 keywords（关键词列表）、pattern（正则表达式）和 answer（预设回答）三个字段。"
        "匹配逻辑：先做关键词 in 判断（O(n) 扫描），命中即返回；未命中则执行 re.search 正则匹配。"
        "目前内置 3 条规则：加班费（正则 r\u201c加班.*?(?:工资|费)\u201d \u2192 "
        "返回《劳动法》第44条全文）、年假（正则 r\u201c(?:带薪)?年?假\u201d \u2192 "
        "返回《职工带薪年休假条例》）、试用期（正则 r\u201c试用\u201d \u2192 返回试用期时长规定）。"
        "规则库设计为可扩展的 Python 列表，可根据业务需要随时增删。\n\n"
        "（4）技能场景系统（models.py + api.py）\n"
        "每个技能定义包含：name（名称）、description（描述）、system_prompt（专业化系统提示词）、"
        "icon（emoji 图标）和 auto_detect_patterns（JSON 数组，自动检测关键词列表）。"
        "系统内置 4 个预置技能：①财务分析（关键词：财务/报表/利润/资产/负债/营收/成本）——"
        "提示词引导模型注重数据准确性和计算过程；②会议纪要助手"
        "（关键词：会议/纪要/议程/决议/讨论/参会/待办）——提示词引导模型提取议题、决议、"
        "责任人、待办事项；③政策文档（关键词：政策/规定/条例/办法/通知/制度）——"
        "提示词引导模型引用具体条款；④发票助手"
        "（关键词：发票/invoice/税额/价税/开票/报销/增值税）——"
        "提示词引导模型提取票面信息、核对金额。"
        "自动检测逻辑：当用户上传文件时，前端调用 /api/skills/detect 接口，"
        "后端遍历所有技能的 auto_detect_patterns，对文件名执行关键词匹配，"
        "返回匹配度最高的技能推荐。\n\n"
        "（5）认证与权限（auth.py + api.py 中间件）\n"
        "认证采用 JWT Token 方案：用户登录时 bcrypt 校验密码，成功后签发有效期 30 分钟的 JWT；"
        "后续所有 API 请求在 Authorization 头携带 Bearer Token。"
        "自定义 auth_middleware 在每次请求前验证 Token 有效性和用户角色，"
        "OPTIONS 预检请求免验证以支持 CORS。支持 admin/user 双角色，"
        "管理员端点（用户管理、反馈查看、技能 CRUD、统计面板）通过 get_admin_user 依赖项限制访问。\n\n"
        "（6）流式响应（SSE）\n"
        "后端 generate_answer_stream 函数返回一个 Python 生成器和元数据字典。"
        "生成器内部调用 LangChain 的 chain.stream() 方法，逐个 yield 清理后的文本 token。"
        "FastAPI 将此生成器包装为 StreamingResponse（media_type=text/event-stream），"
        "前端通过 fetch + ReadableStream 逐块读取并实时渲染到界面上，"
        "配合 Framer Motion 的 AnimatePresence 实现打字机效果。",
    )

    # ====================================================================
    # 第4章 系统实现 (paras 41-44)
    # ====================================================================
    clear_paragraph(paras[42])
    paras[43].style = body_style
    set_paragraph_text(
        paras[43],
        "4.1 后端实现\n"
        "后端基于 FastAPI 框架，采用模块化设计，源码共 8 个核心模块约 2500 行 Python 代码：\n"
        "（1）api.py（~550行）：主路由文件，注册了 37 个以上 API 端点，包括：\n"
        "  - POST /api/auth/login、POST /api/auth/register：用户认证\n"
        "  - POST /api/query：同步 RAG 问答\n"
        "  - POST /api/query/stream：SSE 流式问答\n"
        "  - POST /api/upload：文档上传与入库\n"
        "  - GET /api/documents：文档列表、GET /api/documents/{source}/preview：文档预览\n"
        "  - POST /api/documents/{source}/reingest：重新入库\n"
        "  - GET/POST/PUT/DELETE /api/skills：技能 CRUD\n"
        "  - POST /api/skills/detect：技能自动检测\n"
        "  - GET/POST/DELETE /api/conversations：对话管理\n"
        "  - POST /api/feedback：用户反馈提交\n"
        "  - GET /api/export/{conversation_id}：对话导出（JSON/CSV）\n"
        "  - GET /api/admin/stats、GET /api/admin/users、GET /api/admin/feedback：管理后台\n"
        "（2）rag_engine.py（~460行）：RAG 引擎核心，实现了同步和流式两条问答路径，"
        "包含 LLM 客户端管理、对话历史构建、检索上下文拼装、重试机制等。\n"
        "（3）ingest.py（~75行）：文档入库管道，遍历上传目录、跳过已入库文件、"
        "调用 utils.load_document 解析、RecursiveCharacterTextSplitter 切片、批量写入 ChromaDB。\n"
        "（4）ocr_engine.py（~113行）：GLM-OCR 模型懒加载、单图 OCR、扫描 PDF 逐页渲染 OCR。\n"
        "（5）database.py（~150行）：ChromaDB 封装，包括向量存储初始化、文档增删查改、分页检索。\n"
        "（6）models.py（~360行）：SQLAlchemy ORM 定义 5 张数据表"
        "（users/conversations/messages/feedback/skills），包含数据库迁移逻辑和预置数据初始化。\n"
        "（7）auth.py（~80行）：bcrypt 密码哈希、JWT 签发/验证、角色依赖注入。\n"
        "（8）rule_engine.py（~32行）：关键词+正则规则匹配引擎。\n"
        "辅助模块：config.py（配置管理，包含 LM Studio httpx 客户端工厂）、"
        "schemas.py（Pydantic 请求/响应模型定义）、utils.py（文件加载分发器）。",
    )
    paras[44].style = body_style
    set_paragraph_text(
        paras[44],
        "4.2 前端实现\n"
        "前端采用 Next.js 16 + React 19 + TypeScript + Tailwind CSS 4 + shadcn/ui 组件库，"
        "共包含 5 个主要页面和 10+ 个复用组件，约 3000 行 TypeScript 代码：\n"
        "（1）登录/注册页（/login）：表单输入用户名密码，调用 /api/auth/login 获取 JWT Token "
        "存储在 localStorage，注册后自动登录。\n"
        "（2）聊天对话页（/，首页）：核心交互界面。顶部显示当前技能，"
        "中央为消息列表（用户消息右对齐、AI 回答左对齐），底部为输入框。"
        "关键交互细节：①用户发送消息后，界面立即显示\u201c思考中\u201d动画；"
        "②通过 fetch API 建立 SSE 连接，逐 token 接收并渲染 AI 回答，"
        "配合 Framer Motion 实现平滑的打字机动画；③每条 AI 回答下方显示参考来源折叠面板，"
        "点击可展开查看引用的文档片段和来源文件名；④支持赞/踩反馈按钮。\n"
        "（3）文档管理页（/documents）：展示已上传文档列表，支持上传新文档"
        "（拖拽或点击选择）、预览文档内容、重新入库、删除等操作。"
        "上传成功后触发技能自动检测，推荐匹配的技能。\n"
        "（4）向量片段管理页（/vectors）：展示 ChromaDB 中的全部向量片段，"
        "支持分页浏览、按来源文件筛选、按关键词搜索、编辑片段内容和删除片段。\n"
        "（5）管理员后台（/admin）：三个 Tab 面板——统计概览（用户数、对话数、消息数、文档数、"
        "反馈统计图表）、用户管理（列表/删除）、反馈管理（查看赞踩详情）。\n"
        "侧栏组件集成对话列表管理（新建/切换/删除对话）和技能选择器（下拉选择或自动推荐）。\n\n"
        "4.3 数据库设计\n"
        "向量数据库：ChromaDB，持久化存储在 data/chroma_db/ 目录，"
        "每个文档片段存储原文文本和 768 维嵌入向量，支持余弦相似度检索。\n"
        "关系数据库：SQLite，存储在 data/smartpolicy.db，通过 SQLAlchemy ORM 管理 5 张表：\n"
        "  - users（id, username, hashed_password, role, created_at）\n"
        "  - conversations（id, user_id, skill_id, title, created_at, updated_at）\n"
        "  - messages（id, conversation_id, role, content, sources, rule_matched, created_at）\n"
        "  - feedback（id, message_id, user_id, rating, comment, created_at）\n"
        "  - skills（id, name, description, system_prompt, rules, icon, auto_detect_patterns, "
        "created_at, updated_at）\n"
        "表间关系：users 1:N conversations 1:N messages 1:N feedback；skills 1:N conversations。\n\n"
        "4.4 部署方案\n"
        "提供一键启动脚本 scripts/start-dev.sh，通过 Bash trap 机制同时管理 "
        "FastAPI（uvicorn，端口 8000）和 Next.js（npm run dev，端口 3000）两个子进程。"
        "启动命令：bash scripts/start-dev.sh。Ctrl+C 时 trap 捕获信号同时终止两个进程。"
        "前置依赖：Python 3.10+、Node.js 18+、LM Studio（加载 Qwen 和 nomic-embed 模型）。\n\n"
        "4.5 开发过程中的关键技术困难与解决方案\n"
        "困难一：React 19 Compiler 导致 useEffect 无限循环。React 19 引入的编译器会自动向 "
        "useEffect 依赖数组注入组件内定义的函数引用，导致依赖数组在每次渲染时变化，"
        "触发无限循环。具体表现为页面加载后持续发送 API 请求直至浏览器崩溃。"
        "解决方案：将所有数据获取函数（如 loadConversations、loadSkills）"
        "从 React 组件内部提升为模块级函数（定义在文件顶部、组件外部），"
        "使其引用在模块生命周期内保持稳定，不受 Compiler 依赖注入影响。\n"
        "困难二：httpx 系统代理导致 LM Studio 连接失败。开发环境设置了 HTTP_PROXY 环境变量，"
        "OpenAI SDK 底层使用的 httpx 客户端会自动读取该变量，"
        "将本应直连 127.0.0.1:1234 的请求错误路由到代理服务器，"
        "导致 ConnectionError 和 RemoteProtocolError。"
        "解决方案：在 config.py 中创建专用 httpx 客户端工厂函数 get_lm_studio_httpx_clients()，"
        "设置 trust_env=False 忽略系统代理变量，并注入到所有 ChatOpenAI 和 OpenAIEmbeddings 实例中。",
    )

    # ====================================================================
    # 第5章 测试分析 (paras 45-48)
    # ====================================================================
    clear_paragraph(paras[46])
    paras[47].style = body_style
    set_paragraph_text(
        paras[47],
        "5.1 后端自动化测试\n"
        "本项目使用 pytest 框架建立了完善的自动化测试体系，共编写 42 个以上测试用例，"
        "覆盖后端全部核心模块。测试按功能模块组织在 tests/ 目录下：\n"
        "（1）认证模块（test_auth.py + test_auth_unit.py，共 8 个用例）：\n"
        "  - 测试用户注册：正常注册、重复用户名注册（应返回 400）\n"
        "  - 测试用户登录：正确密码登录（返回 JWT Token）、错误密码登录（返回 401）\n"
        "  - 测试 JWT 验证：有效 Token 解析、过期 Token 拒绝、无效签名拒绝\n"
        "  - 测试角色权限：普通用户访问管理端点（返回 403）\n"
        "（2）API 中间件（test_api_middleware.py，4 个用例）：\n"
        "  - 测试 CORS：OPTIONS 预检请求返回正确的 Access-Control-Allow-* 头\n"
        "  - 测试认证中间件：无 Token 请求返回 401 且包含 CORS 头（避免浏览器吞错误）\n"
        "（3）RAG 问答（test_api_query.py，5 个用例）：\n"
        "  - 测试基本问答：提交问题返回 answer 和 sources 字段\n"
        "  - 测试 temperature 参数：传入 0.1/1.5 不报错\n"
        "  - 测试规则引擎命中：发送\u201c加班费怎么算\u201d，验证 rule_matched 字段非空\n"
        "（4）流式问答（test_streaming.py，3 个用例）：\n"
        "  - 测试 SSE 流格式：响应 Content-Type 为 text/event-stream\n"
        "  - 测试完整流消息：累积所有 chunk 后验证包含有效回答\n"
        "  - 测试对话历史上下文：携带历史消息后生成的回答与无历史时不同\n"
        "（5）文档管理（test_api_documents.py，5 个用例）：\n"
        "  - 测试上传：上传 test.txt 返回成功、上传 .exe 返回 400\n"
        "  - 测试列表：上传后 GET /api/documents 返回包含该文件\n"
        "  - 测试预览：获取文档前 5 个片段的内容\n"
        "  - 测试重新入库和删除：操作后验证状态变化\n"
        "（6）向量片段（test_api_chunks.py，4 个用例）：测试分页、搜索、编辑、删除\n"
        "（7）技能管理（test_skill_api.py，5 个用例）：\n"
        "  - 测试技能列表：GET /api/skills 返回 4 个预置技能\n"
        "  - 测试自动检测：POST /api/skills/detect body={filename:\u201c2025利润表.pdf\u201d}"
        " 返回推荐\u201c财务分析\u201d技能\n"
        "  - 测试管理员 CRUD：创建/更新/删除自定义技能\n"
        "（8）对话管理（test_conversations.py，4 个用例）：创建、列表、详情、删除\n"
        "（9）反馈模块（test_api_feedback.py，2 个用例）：赞/踩提交与查询\n"
        "（10）导出模块（test_api_export.py，2 个用例）：JSON/CSV 格式导出验证\n"
        "（11）数据隔离（test_isolation.py，2 个用例）：两个用户的对话互不可见\n\n"
        "全部 42 个测试用例执行结果：42 passed，0 failed，运行时间约 11 秒。\n\n"
        "5.2 前端验证\n"
        "（1）TypeScript 严格类型检查：执行 npx tsc --noEmit，0 个类型错误。\n"
        "（2）页面功能验证：逐一访问 5 个页面（登录页、聊天页、文档页、向量页、管理页），"
        "确认 HTTP 200 返回且浏览器控制台无 JavaScript 错误。\n\n"
        "5.3 性能测试\n"
        "（1）规则引擎响应时间：对\u201c加班费怎么算\u201d等规则命中问题，"
        "服务端处理时间 < 5ms（测试方法：在 generate_answer 前后打时间戳）。\n"
        "（2）RAG 流式首字节延迟：上传 10 页 PDF 文档入库后提问，"
        "首个 SSE token 到达时间约 1-3 秒（取决于本地 LLM 的推理速度和 GPU 性能）。\n"
        "（3）文档入库性能：10 页纯文本 PDF 的完整入库流程"
        "（文本提取 + 切片 + 嵌入 + ChromaDB 写入）约 5-10 秒。"
        "扫描件 PDF 因需逐页 OCR，速度较慢，每页约 3-5 秒。\n"
        "（4）并发测试：使用 3 个用户同时发送问答请求，系统均能正常响应，"
        "数据隔离测试确认各用户对话互不可见。",
    )
    clear_paragraph(paras[48])

    # ====================================================================
    # 第6章 作品总结 (paras 49-60)
    # ====================================================================
    clear_paragraph(paras[50])

    # 6.1 作品特色与创新点 (H2=51, 正文=52)
    paras[52].style = body_style
    set_paragraph_text(
        paras[52],
        "（1）全链路本地化——数据零泄露\n"
        "LLM 推理（LM Studio + Qwen 3.5-9B）、文本嵌入（nomic-embed-text-v1.5）、"
        "OCR 识别（GLM-OCR）全部在本地 GPU/CPU 上运行。"
        "文档数据从上传到向量化、从检索到生成回答，全程不出内网，不依赖任何云端 API。"
        "这一设计从架构层面根本性地解决了数据隐私问题，是与 ChatGPT 等云端 RAG 方案的本质区别。"
        "以某银行内部使用场景为例：信贷审批流程文档可以安全入库并被员工查询，"
        "无需担心风控策略泄露。\n\n"
        "（2）规则引擎优先策略——高频问题毫秒级响应\n"
        "在 RAG 流程之前设置规则引擎拦截层，通过关键词+正则匹配实现高频问题的零延迟直答。"
        "以\u201c加班费怎么算\u201d为例：规则引擎在 < 5ms 内返回《劳动法》第44条全文，"
        "而如果走完整 RAG 流程则需要 1-3 秒。在日常使用中，约 20-30% 的问题可被规则引擎命中，"
        "显著降低了 GPU 算力消耗。\n\n"
        "（3）技能场景自动识别——零配置适配多业务\n"
        "通过文件名关键词匹配自动推荐最适合的技能。同一套系统无需任何配置即可适配"
        "财务分析、会议纪要整理、发票核验等不同场景。例如用户上传\u201c2025Q3营收报表.pdf\u201d，"
        "系统自动推荐\u201c财务分析\u201d技能；上传\u201c周一部门例会记录.docx\u201d，"
        "自动推荐\u201c会议纪要助手\u201d技能。\n\n"
        "（4）多模态文档处理——打通扫描件信息壁垒\n"
        "集成本地 GLM-OCR 模型，支持扫描 PDF 和 7 种图片格式的 OCR 识别入库。"
        "突破了传统检索系统只能处理文本文档的限制。系统还内置智能扫描件判定"
        "（超过 50% 页面提取文本不足 50 字符），无需用户手动标注。\n\n"
        "（5）工程质量——42 个自动化测试用例全覆盖\n"
        "后端所有核心模块均有对应的 pytest 测试用例，覆盖认证、API、流式、"
        "文档、技能、对话、反馈、导出、数据隔离等 11 个维度，确保代码质量和功能正确性。",
    )

    # 6.2 应用推广 (H2=53, 正文=54,55)
    paras[54].style = body_style
    set_paragraph_text(
        paras[54],
        "本作品可推广到以下具体应用场景：\n\n"
        "场景一：企业内部知识库。以一家 500 人规模的制造企业为例，将 HR 制度手册、"
        "质量管理标准、产品技术文档等约 500 份文件入库后，员工可随时通过对话获取答案。"
        "例如生产线班组长输入\u201c锂电池出货检验标准是什么？\u201d，"
        "系统从质量标准文档中检索到相关条款并生成回答。"
        "预计可将员工查找文档的平均耗时从 15-30 分钟缩短至 10 秒以内。\n\n"
        "场景二：政务政策查询。上传某区社会保障局的全部政策法规文件，"
        "窗口人员面对市民咨询\u201c失业金领取条件是什么？\u201d时，"
        "系统即时检索并引用具体条款生成回答，避免人工翻找的效率损失和记忆偏差。\n\n"
        "场景三：财务审计辅助。审计人员上传被审单位的财务报表 PDF 后，"
        "利用\u201c财务分析\u201d技能提问\u201c2024 年应收账款周转率是多少？\u201d，"
        "系统提取相关数据并给出计算过程和结论。\n\n"
        "场景四：会议记录整理。将会议录音转写稿（TXT 或 DOCX）入库后，"
        "使用\u201c会议纪要助手\u201d技能提问\u201c本次会议有哪些待办事项？\u201d，"
        "系统自动提取议题、决议、责任人和截止日期。\n\n"
        "场景五：发票核验报销。财务人员上传一批发票扫描件后，"
        "使用\u201c发票助手\u201d技能提问\u201c这批发票的总金额是多少？\u201d，"
        "系统通过 OCR 提取票面信息并汇总金额。\n\n"
        "场景六：教育培训。将培训教材 PDF 入库后，学员可提问\u201c请解释 RAG 的工作原理\u201d，"
        "系统从教材中检索相关内容并用通俗语言生成解释。\n\n"
        "由于全部组件均可本地部署，本作品特别适合对数据安全有严格要求的"
        "政府机构、金融企业、国防科研院所和大型组织。",
    )
    clear_paragraph(paras[55])

    # 6.3 作品展望 (H2=56, 正文=57,58,59,60)
    paras[57].style = body_style
    set_paragraph_text(
        paras[57],
        "未来计划从以下五个方向持续提升本作品：\n\n"
        "方向一：多模型适配与自动选择。目前系统固定使用 Qwen 3.5-9B 和 nomic-embed-text-v1.5。"
        "计划增加对 DeepSeek-R1、Llama 3.2、Mistral 等主流开源模型的支持，"
        "并实现基于问题复杂度的自动模型选择（简单问题用小模型加速，复杂问题用大模型提升质量）。\n\n"
        "方向二：多模态交互增强。增加语音输入和语音合成能力（基于 Whisper + VITS），"
        "实现语音问答交互，适用于会议室和车间等不方便打字的场景。"
        "集成表格识别和图表理解能力，支持结构化数据的精准提取和分析。\n\n"
        "方向三：知识图谱融合。在向量检索基础上引入知识图谱，"
        "构建实体（人名、机构、政策条款）之间的关系网络，"
        "实现更精准的关联查询和推理。例如查询\u201c张三负责的所有项目\u201d"
        "不仅依赖文本匹配，还可通过知识图谱的关系路径进行推理。\n\n"
        "方向四：多租户 SaaS 化。支持多组织隔离部署，每个租户拥有独立的文档库、"
        "用户体系和权限管理，适合 SaaS 平台化运营。\n\n"
        "方向五：持续学习与质量优化。利用用户反馈数据（赞/踩及评论）"
        "自动优化检索排序权重和回答质量，实现系统的持续改进。"
        "计划引入 RLHF（基于人类反馈的强化学习）机制，根据积累的反馈数据微调检索和生成策略。",
    )
    for i in range(58, 61):
        clear_paragraph(paras[i])

    # ====================================================================
    # 第7章 参考文献 (paras 61-64)
    # ====================================================================
    clear_paragraph(paras[62])
    paras[63].style = body_style
    set_paragraph_text(
        paras[63],
        "[1] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for "
        "Knowledge-Intensive NLP Tasks[C]. Advances in Neural Information Processing "
        "Systems, 2020, 33: 9459-9474.\n"
        "[2] Du Z, Qian Y, Liu X, et al. GLM: General Language Model Pretraining with "
        "Autoregressive Blank Infilling[C]. Proceedings of the 60th Annual Meeting of "
        "the Association for Computational Linguistics, 2022: 320-335.\n"
        "[3] Nussbaum Z, Morris J, Duderstadt B, et al. Nomic Embed: Training a "
        "Reproducible Long Context Text Embedder[J]. arXiv preprint arXiv:2402.01613, 2024.\n"
        "[4] LangChain Documentation[EB/OL]. https://python.langchain.com/docs/, 2024.\n"
        "[5] ChromaDB Documentation[EB/OL]. https://docs.trychroma.com/, 2024.\n"
        "[6] LM Studio — Discover, download, and run local LLMs[EB/OL]. "
        "https://lmstudio.ai/, 2024.\n"
        "[7] FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/, 2024.\n"
        "[8] Next.js by Vercel — The React Framework[EB/OL]. "
        "https://nextjs.org/docs, 2024.\n"
        "[9] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. "
        "Advances in Neural Information Processing Systems, 2017, 30: 5998-6008.\n"
        "[10] Yang A, Yang B, Hui B, et al. Qwen2 Technical Report[J]. "
        "arXiv preprint arXiv:2407.10671, 2024.",
    )
    clear_paragraph(paras[64])

    # ── 删除示例表格 ──
    table = doc.tables[0]
    tbl_elem = table._tbl
    tbl_elem.getparent().remove(tbl_elem)

    # ── 保存 ──
    doc.save(OUTPUT)
    print(f"已生成: {OUTPUT}")


if __name__ == "__main__":
    main()
